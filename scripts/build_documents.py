"""Build ATS-friendly Word documents from the canonical Markdown sources."""

from __future__ import annotations

import argparse
import logging
import re
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor

LOGGER = logging.getLogger(__name__)
ROOT = Path(__file__).resolve().parents[1]
SOURCES = (ROOT / "resume/executive-resume.md", ROOT / "cv/professional-cv.md")


def add_runs(paragraph: Paragraph, text: str) -> None:
    """Preserve inline bold without introducing text boxes or layout tables."""
    for index, segment in enumerate(re.split(r"\*\*(.*?)\*\*", text)):
        run = paragraph.add_run(segment)
        run.bold = index % 2 == 1


def configure(document: DocumentType, *, resume: bool) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.60)
    section.left_margin = section.right_margin = Inches(0.72)
    section.footer_distance = Inches(0.28)

    # The bundled base template can carry a title border; keep the resume monochrome.
    for border in document.styles.element.xpath(".//w:pBdr"):
        border.getparent().remove(border)

    for name in ("Normal", "Title", "Heading 1", "Heading 2", "List Bullet"):
        style = document.styles[name]
        style.font.name = "Arial"
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.widow_control = True
    normal = document.styles["Normal"]
    normal.font.size = Pt(10.5 if resume else 11)
    normal.paragraph_format.space_after = Pt(5 if resume else 8)
    normal.paragraph_format.line_spacing = 1.10 if resume else 1.16

    title = document.styles["Title"]
    title.font.size = Pt(25)
    title.font.bold = True
    title.paragraph_format.space_after = Pt(8)
    title.paragraph_format.keep_with_next = True
    for name, size in (("Heading 1", 12), ("Heading 2", 11 if resume else 12)):
        style = document.styles[name]
        style.font.size = Pt(size)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(11 if resume else 15)
        style.paragraph_format.space_after = Pt(5 if resume else 7)
        style.paragraph_format.keep_with_next = True
    bullet = document.styles["List Bullet"]
    bullet.font.size = normal.font.size
    bullet.paragraph_format.left_indent = Inches(0.14)
    bullet.paragraph_format.first_line_indent = Inches(-0.14)
    bullet.paragraph_format.space_after = Pt(5 if resume else 8)
    bullet.paragraph_format.line_spacing = normal.paragraph_format.line_spacing

    # Page identity helps recruiters reassemble printed pages; body text remains linear.
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label = "Executive Resume" if resume else "Professional CV"
    run = footer.add_run(f"Pasd Putthapipat | {label} | ")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)


def build(source: Path) -> Path:
    if not source.is_file():
        raise FileNotFoundError(source)
    text = source.read_text(encoding="utf-8")
    if not text.startswith("# "):
        raise ValueError(f"Missing document title: {source}")
    document = Document()
    configure(document, resume=source.stem == "executive-resume")
    page_break_pending = False
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line == "<!-- pagebreak -->":
            page_break_pending = True
            continue
        if line.startswith("<!--"):
            continue
        if line.startswith("# "):
            paragraph = document.add_paragraph(line[2:], "Title")
        elif line.startswith("## "):
            paragraph = document.add_paragraph(line[3:], "Heading 1")
        elif line.startswith("### "):
            paragraph = document.add_paragraph(line[4:], "Heading 2")
        else:
            bullet = line.startswith("- ")
            paragraph = document.add_paragraph(style="List Bullet" if bullet else "Normal")
            add_runs(paragraph, line[2:] if bullet else line)
            if not bullet and line.startswith("**") and line.endswith("**"):
                paragraph.paragraph_format.keep_with_next = True
        if page_break_pending:
            paragraph.paragraph_format.page_break_before = True
            page_break_pending = False
    document.core_properties.author = "Pasd Putthapipat"
    document.core_properties.last_modified_by = "Pasd Putthapipat"
    document.core_properties.title = "Executive Resume" if source.stem == "executive-resume" else "Professional CV"
    document.core_properties.subject = "Data and AI leadership"
    document.core_properties.comments = ""
    output = source.with_suffix(".docx")
    document.save(output)
    LOGGER.info("Built %s", output.relative_to(ROOT))
    return output


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.parse_args()
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    for source in SOURCES:
        build(source)


if __name__ == "__main__":
    main()
